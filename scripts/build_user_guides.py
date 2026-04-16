"""
Build upgraded, professional-grade .docx user guides.

Staged copy — does NOT touch the live repo. Output goes to
HYBRIDRAG_LOCAL_ONLY/staged_guides/ for swap-in after QA clears.

Design standards applied:
  - Calibri 11pt body, Segoe UI Bold headings
  - Navy (#003087) heading accent, #1A1A2E body text
  - Callout boxes (TIP, NOTE, WARNING, CAUTION) with colored borders
  - Score threshold charts via python-docx tables with color coding
  - Analogies inline with every tuning parameter
  - Full glossary, TOC placeholder, footer with version
  - RAGAS how-to with industry thresholds
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

STAGE_DIR = Path(__file__).resolve().parent

# --- Colors ---
NAVY = RGBColor(0, 48, 135)       # #003087
BODY_TEXT = RGBColor(26, 26, 46)   # #1A1A2E
SECONDARY = RGBColor(107, 114, 128)  # #6B7280
LINK_BLUE = RGBColor(0, 120, 212)  # #0078D4
GREEN = RGBColor(22, 163, 74)     # #16A34A
AMBER = RGBColor(217, 119, 6)     # #D97706
RED = RGBColor(220, 38, 38)       # #DC2626
WHITE = RGBColor(255, 255, 255)

# Callout background hex for shading
CALLOUT_COLORS = {
    "TIP":     {"border": "0078D4", "bg": "EFF6FF", "label": "TIP"},
    "NOTE":    {"border": "6B7280", "bg": "F3F4F6", "label": "NOTE"},
    "WARNING": {"border": "D97706", "bg": "FFFBEB", "label": "WARNING"},
    "CAUTION": {"border": "DC2626", "bg": "FEF2F2", "label": "CAUTION"},
    "EXAMPLE": {"border": "16A34A", "bg": "F0FDF4", "label": "EXAMPLE"},
}


def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, side: str, hex_color: str, width: int = 12):
    """Set a single border on a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tc_pr.append(borders)
    border_el = parse_xml(
        f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{width}" '
        f'w:space="0" w:color="{hex_color}"/>'
    )
    existing = borders.find(qn(f"w:{side}"))
    if existing is not None:
        borders.remove(existing)
    borders.append(border_el)


def _setup_styles(doc: Document):
    """Configure modern professional styles."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_TEXT
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3

    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = "Segoe UI"
        hstyle.font.color.rgb = NAVY
        hstyle.font.bold = True
        sizes = {1: 18, 2: 15, 3: 13}
        hstyle.font.size = Pt(sizes[level])
        hstyle.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        hstyle.paragraph_format.space_after = Pt(6)


def _heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    return p


def _bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)


def _callout(doc, callout_type: str, text: str):
    """Add a professional callout box using a single-cell table with colored left border."""
    cfg = CALLOUT_COLORS.get(callout_type, CALLOUT_COLORS["NOTE"])
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    _set_cell_shading(cell, cfg["bg"])
    _set_cell_border(cell, "left", cfg["border"], width=24)
    _set_cell_border(cell, "top", cfg["bg"], width=4)
    _set_cell_border(cell, "bottom", cfg["bg"], width=4)
    _set_cell_border(cell, "right", cfg["bg"], width=4)

    p = cell.paragraphs[0]
    label_run = p.add_run(f"{cfg['label']}:  ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    text_run = p.add_run(text)
    text_run.font.size = Pt(10)
    doc.add_paragraph()  # spacer


def _score_table(doc, title: str, rows: list[list[str]]):
    """Add a color-coded score threshold table."""
    _heading(doc, title, level=3)
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row_data in rows[1:]:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(10)
            # Color code the status column
            lower = text.lower()
            if lower in ("excellent", "active", "pass"):
                run.font.color.rgb = GREEN
                run.bold = True
            elif lower in ("good",):
                run.font.color.rgb = LINK_BLUE
            elif lower in ("needs work", "future"):
                run.font.color.rgb = AMBER
            elif lower in ("red flag", "fail"):
                run.font.color.rgb = RED
                run.bold = True


def _knob_entry(doc, name: str, analogy: str, technical: str):
    """Add a tuning knob entry with analogy and technical explanation."""
    _heading(doc, name, level=3)
    _callout(doc, "EXAMPLE", analogy)
    _para(doc, f"Technical: {technical}")


def _add_cover(doc, title: str, subtitle: str):
    """Add a professional cover page."""
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.bold = True
    run.font.name = "Segoe UI"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(14)
    run2.font.color.rgb = SECONDARY

    for _ in range(4):
        doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Version 2026-04-15  |  Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = SECONDARY

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("HybridRAG V2 — Retrieval-Augmented Generation Platform")
    run4.font.size = Pt(11)
    run4.font.color.rgb = SECONDARY

    for _ in range(2):
        doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Prepared by Jeremy Randall")
    run5.font.size = Pt(12)
    run5.font.color.rgb = BODY_TEXT

    doc.add_page_break()


def _add_version_history(doc):
    """Add a controlled-document version history table."""
    _heading(doc, "Document History", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Rev", "Date", "Author", "Description"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY

    history = [
        ["1.0", "2026-04-15", "J. Randall", "Initial release"],
    ]
    for row_data in history:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(10)

    doc.add_page_break()


def _add_toc(doc):
    """Insert a real Word TOC field that auto-populates when opened."""
    _heading(doc, "Table of Contents", level=1)

    # Insert a TOC field code — Word evaluates this on open or when user
    # right-clicks > "Update Field" / presses F9.  Shows H1-H3.
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_char_separate = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
    )
    run3._r.append(fld_char_separate)

    # Placeholder text shown before first update
    run4 = paragraph.add_run("Right-click here and select 'Update Field' to populate, or press Ctrl+A then F9.")
    run4.font.color.rgb = SECONDARY
    run4.font.size = Pt(10)
    run4.italic = True

    run5 = paragraph.add_run()
    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run5._r.append(fld_char_end)

    doc.add_page_break()


def _add_glossary(doc):
    """Add the full glossary."""
    doc.add_page_break()
    _heading(doc, "Glossary", level=1)

    terms = [
        ("Baseline", "A reference measurement used to track system quality over time. Changes are compared against the baseline to detect improvements or regressions."),
        ("Chunk", "A piece of a document, typically 200-500 words, stored as a unit in the vector index. Documents are split into chunks before embedding."),
        ("Context Precision", "RAGAS metric measuring whether the most relevant chunks appear at the top of retrieval results. High precision = low noise in results."),
        ("Context Recall", "RAGAS metric measuring whether all information needed to answer correctly was retrieved. High recall = no important information missed."),
        ("Corpus", "The full collection of documents being searched. Our production corpus contains the ingested document library."),
        ("Embedding", "Converting text into a list of numbers (a vector) that captures its meaning. Similar texts produce similar vectors, enabling semantic search."),
        ("Faithfulness", "RAGAS metric measuring whether every claim in the generated answer is supported by the retrieved context. Prevents hallucination."),
        ("Ground Truth", "The known-correct answer and relevant contexts used to measure system accuracy. Created by human annotation or curated reference data."),
        ("Hallucination", "When the AI generates information that isn't supported by any retrieved document. Dangerous in regulated or factual applications."),
        ("LanceDB", "The vector database that stores our document embeddings and enables fast similarity search."),
        ("Latency", "How long a query takes from submission to response. Measured in milliseconds. p50 = median, p95 = 95th percentile."),
        ("nprobes", "How many sections (partitions) of the vector index to search. More = more thorough but slower."),
        ("Precision", "Of everything retrieved, what fraction was actually relevant."),
        ("RAGAS", "Retrieval Augmented Generation Assessment — the industry-standard open-source framework for evaluating RAG system quality."),
        ("RAG", "Retrieval-Augmented Generation — an AI architecture that looks up relevant documents before generating an answer, grounding responses in real data."),
        ("Recall", "Of everything relevant in the corpus, what fraction was successfully retrieved."),
        ("Regression", "When a code change causes something that previously worked correctly to start failing."),
        ("Reranker", "A second-pass model that re-reads retrieved documents alongside the query and re-orders them by true relevance. Significantly improves precision."),
        ("temperature", "Controls how creative vs deterministic the AI's response is. Low = factual and consistent. High = varied and creative."),
        ("top_k", "In retrieval: how many results to return. In generation: how many candidate words to consider at each step."),
        ("top_p", "Nucleus sampling — only consider candidate words whose cumulative probability reaches this threshold. Keeps responses focused."),
        ("Vector Search", "Finding documents by mathematical similarity of their embeddings to the query embedding, rather than keyword matching."),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.7)
    h_cells = table.rows[0].cells
    for i, h in enumerate(["Term", "Definition"]):
        h_cells[i].text = ""
        run = h_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY

    for term, defn in terms:
        row = table.add_row()
        tc = row.cells[0]
        tc.text = ""
        r = tc.paragraphs[0].add_run(term)
        r.bold = True
        r.font.size = Pt(10)
        dc = row.cells[1]
        dc.text = ""
        r2 = dc.paragraphs[0].add_run(defn)
        r2.font.size = Pt(10)


def _add_ragas_section(doc):
    """Comprehensive RAGAS explainer with industry context."""
    doc.add_page_break()
    _heading(doc, "RAGAS: Industry-Standard RAG Evaluation", level=1)

    _heading(doc, "What is RAGAS?", level=2)
    _para(doc, (
        "RAGAS (Retrieval Augmented Generation Assessment) is the industry-standard "
        "open-source framework for evaluating Retrieval-Augmented Generation systems. "
        "Published as a peer-reviewed research paper (arXiv:2309.15217), RAGAS provides "
        "objective, reproducible metrics that measure both retrieval quality and generation "
        "quality. It is the most widely adopted RAG evaluation framework in production use."
    ))
    _callout(doc, "NOTE", (
        "RAGAS is to RAG systems what unit tests are to code: an objective, repeatable "
        "way to measure whether the system is working correctly and catch regressions "
        "before they reach production."
    ))

    _heading(doc, "The Five Core Industry-Standard Metrics", level=2)
    _para(doc, (
        "A complete RAG evaluation covers five core metrics across two dimensions: "
        "retrieval quality (did we find the right information?) and generation quality "
        "(did we produce a good answer from it?)."
    ))

    _score_table(doc, "Core Metrics Reference", [
        ["Metric", "Dimension", "What It Measures", "Target", "Red Flag"],
        ["Context Precision", "Retrieval", "Are relevant chunks ranked highest?", "0.70+", "Below 0.50"],
        ["Context Recall", "Retrieval", "Did we retrieve all needed info?", "0.75+", "Below 0.50"],
        ["Faithfulness", "Generation", "Are claims supported by context?", "0.80+", "Below 0.70"],
        ["Answer Relevancy", "Generation", "Does the answer address the question?", "0.75+", "Below 0.60"],
        ["Hallucination Rate", "End-to-End", "Unsupported claims in responses", "Below 0.05", "Above 0.10"],
    ])

    # Metric analogies
    _heading(doc, "Understanding Each Metric", level=2)

    _heading(doc, "Context Precision — Signal vs Noise", level=3)
    _callout(doc, "EXAMPLE", (
        "Imagine searching Google and getting 10 results. Context precision asks: "
        "'Were the top results actually relevant, or did junk float to the top?' "
        "A precision of 0.8 means 8 out of 10 retrieved chunks were useful."
    ))
    _para(doc, "Technical: Ratio of relevant retrieved contexts to total retrieved contexts, weighted by rank position. Uses reference contexts for comparison.")

    _heading(doc, "Context Recall — Coverage Check", level=3)
    _callout(doc, "EXAMPLE", (
        "Like a detective checking if they've gathered ALL the evidence before presenting "
        "a case. Missing one key piece means the answer might be wrong or incomplete. "
        "A recall of 0.9 means we found 90% of all relevant information."
    ))
    _para(doc, "Technical: Fraction of reference context statements that appear (by fuzzy match) in the retrieved contexts.")

    _heading(doc, "Faithfulness — Cite Your Sources", level=3)
    _callout(doc, "EXAMPLE", (
        "Like a journalist — every claim in the article should trace back to a named source. "
        "If the AI says 'the system was deployed in March' but no retrieved document mentions "
        "March, that's an unfaithful claim."
    ))
    _para(doc, "Technical: Ratio of claims in the generated answer that are supported by the retrieved context to total claims. Requires an LLM judge (future enhancement).")

    _heading(doc, "Answer Relevancy — Did You Answer the Question?", level=3)
    _callout(doc, "EXAMPLE", (
        "The system might retrieve great context and ground its answer faithfully — but "
        "did it actually answer what was asked? Like a student writing a well-researched "
        "essay on the wrong topic."
    ))
    _para(doc, "Technical: Measures semantic similarity between the generated answer and the original question. Requires an LLM judge (future enhancement).")

    _heading(doc, "Hallucination Rate — Made-Up Facts", level=3)
    _callout(doc, "WARNING", (
        "In medical, legal, or enterprise contexts, even a 5% hallucination rate is dangerous. "
        "This metric tracks the proportion of responses containing claims not backed by "
        "any retrieved evidence. Target: below 2% for regulated industries."
    ))

    # Our implementation mapping
    _heading(doc, "How Our System Maps to Industry Standards", level=2)
    _score_table(doc, "Implementation Status", [
        ["Industry Metric", "Our Implementation", "Status", "Notes"],
        ["Context Precision", "NonLLMContextPrecisionWithReference", "ACTIVE", "Levenshtein similarity via rapidfuzz. Offline, no API."],
        ["Context Recall", "NonLLMContextRecall", "ACTIVE", "Fuzzy match retrieved vs reference contexts. Offline."],
        ["Faithfulness", "Planned (LLM judge)", "FUTURE", "Requires external API. Measures claim grounding."],
        ["Answer Relevancy", "Planned (LLM judge)", "FUTURE", "Requires external API. Measures Q-A alignment."],
        ["Hallucination Rate", "Derived from Faithfulness", "FUTURE", "Computed as 1.0 - Faithfulness."],
        ["Precision@K / Recall@K", "Production eval scorer", "ACTIVE", "PASS/PARTIAL/MISS verdicts in 400-query pack."],
        ["Per-query latency", "Production eval runner", "ACTIVE", "p50/p95 wall-clock, router, embed, retrieval."],
    ])

    _callout(doc, "TIP", (
        "LLM vs Non-LLM metrics: Our current metrics use algorithmic fuzzy matching "
        "(no external API needed). LLM-based metrics use a judge model for deeper semantic "
        "evaluation but require API access and cost. The Non-LLM variants are the correct "
        "choice for offline, offline, and cost-controlled environments."
    ))

    # RAGAS How-To
    _heading(doc, "Running RAGAS Evaluations — How-To", level=2)

    _heading(doc, "What You Need", level=3)
    _bullet(doc, "", bold_prefix="Query pack: ")
    _para(doc, "  JSON file with ground-truth annotations (reference answers + reference contexts). Our 400-query production pack has this via Phase 2C enrichment.")
    _bullet(doc, "", bold_prefix="Dependencies: ")
    _para(doc, "  ragas and rapidfuzz installed in the .venv (run tools/install_ragas_proxy_ready_2026-04-15.ps1 if needed)")
    _bullet(doc, "", bold_prefix="Live index: ")
    _para(doc, "  LanceDB populated with embedded corpus")

    _heading(doc, "Quick Run (Analysis Only)", level=3)
    _para(doc, "1. Open QA Workbench or Eval GUI")
    _para(doc, "2. Click the RAGAS tab")
    _para(doc, '3. Check "Analysis only" checkbox')
    _para(doc, "4. Click Start")
    _para(doc, "5. Review: eligible count, phase2c count, dependency status")
    _callout(doc, "TIP", "Analysis-only mode completes in seconds and needs no GPU. Use it to verify your query pack and dependencies before a full run.")

    _heading(doc, "Full Execution Run", level=3)
    _para(doc, '1. Uncheck "Analysis only"')
    _para(doc, "2. Set Limit to 3-10 for a quick test, or leave blank for all eligible queries")
    _para(doc, "3. Click Start")
    _para(doc, "4. Watch the live log — each query shows retrieved context count")
    _para(doc, "5. When done, review metric means in the summary area")
    _para(doc, "6. Click Open Artifact to see the full JSON results")

    _heading(doc, "Score Interpretation", level=3)
    _score_table(doc, "What Your Scores Mean", [
        ["Score Range", "Rating", "What To Do"],
        ["0.90+", "Excellent", "System is performing at top tier. Document and maintain."],
        ["0.70 - 0.89", "Good", "Production-ready. Look for targeted improvements."],
        ["0.50 - 0.69", "Needs Work", "Check chunk quality, embedding model, nprobes, or reranker."],
        ["Below 0.50", "Red Flag", "Significant issues. Investigate corpus coverage and embedding alignment."],
    ])

    _callout(doc, "WARNING", (
        "A meaningful improvement is 0.05+ on any single metric. Changes smaller than 0.02 "
        "may be noise. Always compare on the same query pack and same corpus version — one "
        "variable at a time."
    ))


def _add_tuning_knobs(doc):
    """Full tuning knobs reference with analogies."""
    doc.add_page_break()
    _heading(doc, "Tuning Knobs Reference", level=1)
    _para(doc, "Every adjustable parameter for retrieval and generation quality, explained in plain language and technical detail.")

    _heading(doc, "Retrieval Parameters — Finding the Right Information", level=2)

    _knob_entry(doc, "top_k — Retrieval Depth",
        "Like asking a store clerk to bring you their top 5 vs top 20 best matches. "
        "More picks = more likely to find the right one, but also more irrelevant items to sort through.",
        "Number of nearest-neighbor vectors returned by embedding search. "
        "Higher values increase recall (fewer misses) but decrease precision (more noise). "
        "Default: 5. Range: 1-50. Note: top_k is currently shadowed by TOP_K=5 in the production eval path.")

    _knob_entry(doc, "candidate_pool — Search Width",
        "Like screening 100 resumes to pick 5 finalists for interviews. The wider "
        "your initial pool, the less likely you'll miss a great candidate hiding in the stack.",
        "Pre-filter pool size before final ranking. Larger pools catch edge-case matches "
        "at higher compute cost. Typical: 50-200. Config: retrieval.candidate_pool in config.yaml.")

    _knob_entry(doc, "nprobes — Index Thoroughness",
        "Your index is organized into sections, like aisles in a library. "
        "nprobes controls how many aisles the search checks. More aisles = more thorough but slower.",
        "Number of IVF partitions scanned in the LanceDB vector index. "
        "Higher nprobes = better recall at cost of latency. Set via HYBRIDRAG_LANCE_NPROBES env var.")

    _knob_entry(doc, "refine_factor — Result Refinement",
        "Like pulling 3x the books from the filing cabinet, then carefully re-reading each "
        "one with fresh eyes to pick the truly best matches.",
        "LanceDB refine factor. Fetches refine_factor * nprobes candidates, re-ranks with "
        "full-precision vectors. Higher = better accuracy, more compute. Set via HYBRIDRAG_LANCE_REFINE_FACTOR.")

    _knob_entry(doc, "reranker_enabled — Two-Stage Ranking",
        "The fast initial search finds rough matches quickly. The reranker is like calling "
        "in a specialist who carefully re-reads each candidate alongside the original question "
        "and re-ranks them based on true relevance.",
        "When enabled, a cross-encoder (FlashRank or similar) rescores candidate_pool using "
        "full query-document pairs. Significantly improves precision. Config: retrieval.reranker_enabled.")

    _heading(doc, "Generation Parameters — Producing the Answer", level=2)

    _knob_entry(doc, "temperature — Creativity vs Consistency",
        "Like a thermostat for creativity. Low = the system gives the most predictable, "
        "factual answer every time. High = more varied and creative, but higher risk of "
        "making things up. For factual QA: keep it cold (0.0-0.3).",
        "Controls softmax distribution over next-token probabilities. "
        "0.0 = deterministic (greedy). >0.5 = increasing randomness. Range: 0.0-2.0.")

    _knob_entry(doc, "top_p — Nucleus Sampling",
        "Imagine all possible next words standing in a crowd, sorted by how likely they are. "
        "top_p says 'only let the people in the brightest 90% of the spotlight speak.' "
        "Keeps responses focused and on-topic.",
        "Cumulative probability threshold for token selection. "
        "top_p=0.9 means only the top 90% probability mass is considered. Range: 0.0-1.0.")

    _knob_entry(doc, "top_k (Generation) — Word Candidate Limit",
        "Instead of considering every word in the dictionary for the next word, "
        "only consider the K most likely candidates. Like limiting a ballot to "
        "serious candidates only.",
        "Limits token selection to K highest-probability candidates. top_k=50 is common. "
        "Use with temperature > 0. Not applicable when temperature = 0 (greedy decoding).")

    _knob_entry(doc, "max_tokens — Response Length Limit",
        "Like a word count limit on an essay. Prevents the system from rambling, "
        "but set it too low and the answer gets cut off mid-sentence.",
        "Hard cap on generated output tokens. Too low truncates; too high wastes on padding. "
        "Typical: 256-1024 for QA.")

    _knob_entry(doc, "repetition_penalty — Anti-Repetition",
        "Discourages the model from repeating the same phrases over and over, "
        "like a writing coach saying 'find a different way to say that.'",
        "Multiplicative penalty on already-generated tokens. 1.0 = none. "
        "1.1-1.3 = mild. Higher may break coherence.")

    _knob_entry(doc, "context_window — Input Capacity",
        "How many retrieved documents the model can spread out and read at once. "
        "Like the size of your desk — a bigger desk lets you reference more materials "
        "simultaneously, but costs more to maintain.",
        "Maximum input token length for the generation model. Determines how many "
        "retrieved chunks fit in a single prompt. Models range from 4K to 128K+ tokens.")


def _add_references(doc):
    """Add a references section with proper citations."""
    doc.add_page_break()
    _heading(doc, "References", level=1)

    refs = [
        ("RAGAS Framework",
         'Es, S., James, J., Espinosa-Anke, L., & Schockaert, S. (2023). '
         '"RAGAS: Automated Evaluation of Retrieval Augmented Generation." '
         'arXiv:2309.15217. https://arxiv.org/abs/2309.15217'),
        ("RAGAS Documentation",
         'Explodinggradients. RAGAS — Retrieval Augmented Generation Assessment. '
         'https://docs.ragas.io'),
        ("Score Thresholds",
         'Industry benchmark thresholds adapted from Qdrant RAG Evaluation Guide '
         '(qdrant.tech/blog/rag-evaluation-guide) and PremAI RAG Evaluation '
         'Metrics & Testing 2026 (blog.premai.io).'),
        ("Microsoft Writing Style Guide",
         'Microsoft Corporation. Microsoft Writing Style Guide. '
         'https://learn.microsoft.com/en-us/style-guide/welcome/'),
        ("Typography Standards",
         'Document typography follows recommendations from the Google Developer '
         'Documentation Style Guide and ISO/IEC Directives Part 2 for '
         'structure and drafting of technical documents.'),
    ]

    for title, citation in refs:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{title}:  ")
        run_t.bold = True
        run_t.font.size = Pt(10)
        run_c = p.add_run(citation)
        run_c.font.size = Pt(10)
        run_c.font.color.rgb = SECONDARY


def _add_footer(doc, text: str):
    """Add footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = SECONDARY


def _build_qa_workbench_guide():
    doc = Document()
    _setup_styles(doc)

    _add_cover(doc, "QA Workbench", "User Guide")
    _add_toc(doc)
    _add_version_history(doc)

    # Quick Start
    _heading(doc, "Quick Start", level=1)
    _para(doc, "Get the QA Workbench running in under 5 minutes.")
    _heading(doc, "Prerequisites", level=2)
    _bullet(doc, "Windows 10/11 with NVIDIA GPU (CUDA required)")
    _bullet(doc, "HybridRAG V2 repo cloned with .venv installed (run INSTALL_EVAL_GUI.bat if needed)")
    _bullet(doc, "LanceDB index built and populated (data/ directory)")
    _heading(doc, "Launch", level=2)
    _para(doc, "Double-click start_qa_workbench.bat, or from a terminal:")
    _para(doc, "cd /d C:\\HybridRAG_V2 && start_qa_workbench.bat", italic=True)
    _callout(doc, "TIP", "Behind a corporate proxy? Set HTTPS_PROXY in your shell before launching. The launcher inherits proxy settings automatically.")
    _callout(doc, "NOTE", "The window has a vertical scrollbar. You can shrink it to half-screen and scroll to reach all content.")

    # Tab guides
    _heading(doc, "Tab 1: Overview", level=1)
    _para(doc, "A read-only management dashboard showing all measurement lanes at a glance.")
    _bullet(doc, "Certified baseline status and scores")
    _bullet(doc, "Provider comparison (speed and quality)")
    _bullet(doc, "Count, Aggregation, and Regression benchmark status")
    _bullet(doc, "Strongest/weakest persona and query type breakdown")
    _callout(doc, "NOTE", "No setup needed. This tab reads existing result files. Missing data shows '(not yet available)' gracefully.")

    _heading(doc, "Tab 2: Baseline (Production Eval)", level=1)
    _para(doc, "Run the full 400-query production evaluation against the live index. This is the primary quality measurement tool.")
    _heading(doc, "Setup", level=2)
    _bullet(doc, "Query pack defaults to production_queries_400_2026-04-12.json (400 curated queries, 3 personas)")
    _bullet(doc, "Config defaults to config.yaml — change only for specific config testing")
    _bullet(doc, "GPU index: check nvidia-smi to see which GPU is available")
    _bullet(doc, "Max Queries: blank for full 400, or 5-10 for a smoke test")
    _heading(doc, "Running", level=2)
    _para(doc, "Click Start. Progress bar advances per query. Live log shows color-coded verdicts:")
    _bullet(doc, "Green = PASS (expected answer found in retrieved context)")
    _bullet(doc, "Orange = PARTIAL (some relevant content, incomplete)")
    _bullet(doc, "Red = MISS (expected content not retrieved)")
    _para(doc, "Click Stop at any time to safely halt after the current query.")
    _heading(doc, "Interpreting Results", level=2)
    _callout(doc, "WARNING", "PASS rate below 60% indicates a significant retrieval problem. Investigate corpus coverage and embedding quality before proceeding.")
    _bullet(doc, "p95 latency above 30s suggests a performance issue — check GPU load")
    _bullet(doc, "High MISS rate on one persona often indicates a corpus gap for that query type")
    _para(doc, "Sub-tabs: Results (browse), Compare (diff two runs), History (all runs over time)")

    _heading(doc, "Tab 3: Aggregation", level=1)
    _para(doc, "Tests cross-document aggregation accuracy. Can the system combine information from multiple sources correctly?")
    _bullet(doc, "Manifest: 12-item aggregation seed (default)")
    _bullet(doc, "Self-check mode (blank answers) is fastest — tests against certified baseline")
    _bullet(doc, "12/12 pass = baseline intact. Any failure = regression.")
    _callout(doc, "TIP", "Run Aggregation after any retrieval change. It catches cross-document combination errors that single-query tests miss.")

    _heading(doc, "Tab 4: Count", level=1)
    _para(doc, "Verifies corpus coverage by counting entity mentions, documents, chunks, and rows.")
    _bullet(doc, "7/7 frozen-expectation match = corpus complete, dedup stable")
    _bullet(doc, "Count mismatch = new data ingested (expected) or dedup changed (investigate)")
    _callout(doc, "NOTE", "Count benchmark requires LanceDB and Entity DB paths. The defaults point to the standard install locations.")

    _heading(doc, "Tab 5: RAGAS", level=1)
    _para(doc, "Industry-standard retrieval quality measurement. See the RAGAS Reference section for full details.")
    _bullet(doc, "Analysis Only: fast readiness check, no GPU needed")
    _bullet(doc, "Full Execution: measures context precision and context recall per query")
    _bullet(doc, "Set Limit to 3-10 for quick tests")

    _heading(doc, "Tab 6: Regression", level=1)
    _para(doc, "Fast guardrail — 50 frozen extraction patterns across 5 entity families.")
    _bullet(doc, "Completes in 1-3 seconds, no database or GPU needed")
    _bullet(doc, "50/50 PASS = all known patterns work")
    _bullet(doc, "Any failure = a code change broke an extraction pattern")

    _heading(doc, "Tab 7: History / Ledger", level=1)
    _para(doc, "Browse all past evaluation runs in a sortable table. Track quality over time.")
    _callout(doc, "TIP", "After any retrieval-side code change, the recommended benchmark order is: Regression > Count > Aggregation > Baseline > append ledger row.")

    # RAGAS deep-dive
    _add_ragas_section(doc)
    _add_tuning_knobs(doc)
    _add_glossary(doc)

    # Troubleshooting
    doc.add_page_break()
    _heading(doc, "Troubleshooting", level=1)
    _callout(doc, "CAUTION", "Window opens and closes immediately? Run from terminal to see the error. Most common: missing .venv or broken Python. Run INSTALL_EVAL_GUI.bat.")
    _bullet(doc, "CUDA not available: Check nvidia-smi. Set CUDA_VISIBLE_DEVICES=0 before launching.")
    _bullet(doc, "Proxy errors: Set HTTPS_PROXY and HTTP_PROXY in your shell before launch.")
    _bullet(doc, "RAGAS blocked: Run tools/install_ragas_proxy_ready_2026-04-15.ps1")
    _bullet(doc, "Results not in History: Results must be saved as production_eval_results*.json in docs/")

    _add_references(doc)
    _add_footer(doc, "HybridRAG V2 — QA Workbench User Guide — 2026-04-15 — J. Randall")

    path = STAGE_DIR / "QA_WORKBENCH_USER_GUIDE_v2.docx"
    doc.save(str(path))
    print(f"Written: {path}")


def _build_eval_gui_guide():
    doc = Document()
    _setup_styles(doc)

    _add_cover(doc, "Eval GUI", "User Guide")
    _add_toc(doc)

    _heading(doc, "Quick Start", level=1)
    _bullet(doc, "Windows 10/11, NVIDIA GPU, .venv installed")
    _bullet(doc, "Launch: start_eval_gui.bat")
    _bullet(doc, "8 tabs: Overview, Launch, Aggregation, Count, RAGAS, Results, Compare, History")
    _callout(doc, "TIP", "The Eval GUI is the operator's tool. The QA Workbench adds management-facing summary strips and regression. Both share the same core panels.")

    _heading(doc, "Tab: Launch", level=1)
    _para(doc, "Run the 400-query production evaluation. Same functionality as QA Workbench Baseline tab.")
    _bullet(doc, "Select query pack, config, GPU index, optional limit")
    _bullet(doc, "Start/Stop with live progress and color-coded verdicts")
    _bullet(doc, "Save as Defaults persists your input paths")

    _heading(doc, "Tab: Results", level=1)
    _para(doc, "Browse and filter a completed evaluation run in detail.")
    _bullet(doc, "Filter by Verdict, Persona, Family, or Query Type")
    _bullet(doc, "Click a row for full details: routing, top chunks, per-stage timing")
    _callout(doc, "WARNING", "Watch for routing mismatches (expected vs actual query type) and empty retrieved chunks (retrieval failures).")

    _heading(doc, "Tab: Compare", level=1)
    _para(doc, "Diff two evaluation runs side-by-side.")
    _bullet(doc, "Select Run A (before) and Run B (after)")
    _bullet(doc, "Green rows = improvements, Red = regressions")
    _bullet(doc, "Filter: Changed Only, Gains Only, Losses Only")
    _callout(doc, "TIP", "Any PASS > MISS transition warrants investigation before accepting the code change.")

    _heading(doc, "Tabs: Aggregation, Count, RAGAS, History", level=1)
    _para(doc, "Same panels as QA Workbench. See the QA Workbench User Guide for detailed instructions.")

    _add_ragas_section(doc)
    _add_tuning_knobs(doc)
    _add_glossary(doc)

    doc.add_page_break()
    _heading(doc, "Troubleshooting", level=1)
    _bullet(doc, "Same troubleshooting steps as QA Workbench. See QA Workbench User Guide.")

    _add_references(doc)
    _add_footer(doc, "HybridRAG V2 — Eval GUI User Guide — 2026-04-15 — J. Randall")

    path = STAGE_DIR / "EVAL_GUI_USER_GUIDE_v2.docx"
    doc.save(str(path))
    print(f"Written: {path}")


def main():
    _build_qa_workbench_guide()
    _build_eval_gui_guide()
    print("\nStaged guides ready for swap-in after QA clears.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
