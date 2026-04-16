"""Convert all markdown docs to .docx format using python-docx."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx(md_path: Path, docx_path: Path):
    """Convert a markdown file to .docx with basic formatting."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    text = md_path.read_text(encoding='utf-8-sig')
    lines = text.split('\n')
    in_code_block = False
    code_buffer = []

    for line in lines:
        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - flush buffer
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(40, 40, 40)
                p.paragraph_format.left_indent = Inches(0.3)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        stripped = line.strip()

        # Headers
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)

        # Horizontal rule
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

        # Table rows (simple rendering)
        elif stripped.startswith('|') and stripped.endswith('|'):
            # Skip separator rows
            if all(c in '|-: ' for c in stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            cell_text = '  |  '.join(cells)
            p = doc.add_paragraph(cell_text)
            run = p.runs[0] if p.runs else p.add_run('')
            run.font.size = Pt(10)
            run.font.name = 'Consolas'

        # Checklist items
        elif stripped.startswith('- [ ] '):
            doc.add_paragraph(stripped[6:], style='List Bullet')
        elif stripped.startswith('- [x] '):
            doc.add_paragraph(stripped[6:] + ' (DONE)', style='List Bullet')

        # Bullet points
        elif stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')

        # Bold lines (metadata)
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.bold = True

        # Empty lines
        elif not stripped:
            doc.add_paragraph()

        # Normal text
        else:
            # Handle inline bold
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            clean = re.sub(r'`(.+?)`', r'\1', clean)
            clean = re.sub(r'~~(.+?)~~', r'\1', clean)
            doc.add_paragraph(clean)

    doc.save(str(docx_path))
    print(f"  Created: {docx_path.name}")


def main():
    """Parse command-line inputs and run the main convert docs to docx workflow."""
    docs_dir = Path(__file__).parent.parent / 'docs'
    md_files = sorted(docs_dir.glob('*.md'))

    print(f"Converting {len(md_files)} markdown files to .docx...")
    for md_file in md_files:
        docx_file = md_file.with_suffix('.docx')
        md_to_docx(md_file, docx_file)

    print("Done.")


if __name__ == '__main__':
    main()
